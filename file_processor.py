import os
import json
import pandas as pd
from docx import Document
from pypdf import PdfReader
from typing import List, Dict, Any
import tempfile

class FileProcessor:
    """Process various file types and extract text content"""
    
    @staticmethod
    def process_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract content
        
        Args:
            file_path: Path to the uploaded file
            file_type: Type of file (pdf, docx, xlsx, csv, txt, json)
            
        Returns:
            Dictionary with extracted content and metadata
        """
        try:
            if file_type == 'pdf':
                return FileProcessor._process_pdf(file_path)
            elif file_type == 'docx':
                return FileProcessor._process_docx(file_path)
            elif file_type == 'xlsx':
                return FileProcessor._process_xlsx(file_path)
            elif file_type == 'csv':
                return FileProcessor._process_csv(file_path)
            elif file_type == 'txt':
                return FileProcessor._process_txt(file_path)
            elif file_type == 'json':
                return FileProcessor._process_json(file_path)
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_type}"
                }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error processing file: {str(e)}"
            }
    
    @staticmethod
    def _process_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return {
            "status": "success",
            "file_type": "pdf",
            "content": text.strip(),
            "metadata": {
                "num_pages": len(reader.pages),
                "total_characters": len(text)
            }
        }
    
    @staticmethod
    def _process_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            "status": "success",
            "file_type": "docx",
            "content": text.strip(),
            "metadata": {
                "num_paragraphs": len(doc.paragraphs),
                "total_characters": len(text)
            }
        }
    
    @staticmethod
    def _process_xlsx(file_path: str) -> Dict[str, Any]:
        """Extract data from Excel"""
        df = pd.read_excel(file_path)
        text = df.to_string()
        
        return {
            "status": "success",
            "file_type": "xlsx",
            "content": text,
            "data": df.to_dict(orient='records'),
            "metadata": {
                "num_rows": len(df),
                "num_columns": len(df.columns),
                "columns": list(df.columns)
            }
        }
    
    @staticmethod
    def _process_csv(file_path: str) -> Dict[str, Any]:
        """Extract data from CSV"""
        df = pd.read_csv(file_path)
        text = df.to_string()
        
        return {
            "status": "success",
            "file_type": "csv",
            "content": text,
            "data": df.to_dict(orient='records'),
            "metadata": {
                "num_rows": len(df),
                "num_columns": len(df.columns),
                "columns": list(df.columns)
            }
        }
    
    @staticmethod
    def _process_txt(file_path: str) -> Dict[str, Any]:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            "status": "success",
            "file_type": "txt",
            "content": text.strip(),
            "metadata": {
                "total_characters": len(text),
                "num_lines": len(text.split('\n'))
            }
        }
    
    @staticmethod
    def _process_json(file_path: str) -> Dict[str, Any]:
        """Extract data from JSON"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        text = json.dumps(data, indent=2)
        
        return {
            "status": "success",
            "file_type": "json",
            "content": text,
            "data": data,
            "metadata": {
                "total_characters": len(text)
            }
        }
